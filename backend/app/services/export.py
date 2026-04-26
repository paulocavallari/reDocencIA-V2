from __future__ import annotations

import io
import re
from html import unescape

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


def _extract_text_with_format(node: Tag | NavigableString, paragraph, bold: bool = False, italic: bool = False, underline: bool = False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text.strip():
            run = paragraph.add_run(unescape(text))
            run.bold = bold
            run.italic = italic
            run.underline = underline
        return

    next_bold = bold or node.name in {"strong", "b"}
    next_italic = italic or node.name in {"em", "i"}
    next_underline = underline or node.name == "u"
    for child in node.children:
        _extract_text_with_format(child, paragraph, next_bold, next_italic, next_underline)


def _alignment_from_tag(node: Tag) -> WD_ALIGN_PARAGRAPH | None:
    class_attr = " ".join(node.get("class", []))
    style_attr = node.get("style", "")
    if "ql-align-center" in class_attr or "text-align: center" in style_attr:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "ql-align-right" in class_attr or "text-align: right" in style_attr:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "ql-align-justify" in class_attr or "text-align: justify" in style_attr:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _clean_html_text(text: str) -> str:
    normalized = re.sub(r"\s+", " ", unescape(text or "")).strip()
    return normalized


def build_docx(plan_title: str, html: str) -> io.BytesIO:
    document = Document()
    title = document.add_heading(plan_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    soup = BeautifulSoup(html, "html.parser")
    root_nodes = soup.body.contents if soup.body else soup.contents

    for node in root_nodes:
        if isinstance(node, NavigableString):
            if _clean_html_text(str(node)):
                document.add_paragraph(_clean_html_text(str(node)))
            continue

        if node.name in {"h1", "h2", "h3"}:
            level = {"h1": 1, "h2": 2, "h3": 3}[node.name]
            document.add_heading(_clean_html_text(node.get_text(" ")), level=level)
            continue

        if node.name in {"ul", "ol"}:
            style = "List Bullet" if node.name == "ul" else "List Number"
            for item in node.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style=style)
                _extract_text_with_format(item, paragraph)
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = _alignment_from_tag(node)
        _extract_text_with_format(node, paragraph)
        if not paragraph.text.strip():
            document._body._element.remove(paragraph._element)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _html_to_reportlab_markup(node: Tag | NavigableString) -> str:
    if isinstance(node, NavigableString):
        return unescape(str(node)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    child_markup = "".join(_html_to_reportlab_markup(child) for child in node.children)
    if node.name in {"strong", "b"}:
        return f"<b>{child_markup}</b>"
    if node.name in {"em", "i"}:
        return f"<i>{child_markup}</i>"
    if node.name == "u":
        return f"<u>{child_markup}</u>"
    if node.name == "br":
        return "<br/>"
    return child_markup


def build_pdf(plan_title: str, html: str) -> io.BytesIO:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    title_style.textColor = colors.HexColor("#163344")
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
        textColor=colors.HexColor("#163344"),
    )
    bullet_style = ParagraphStyle("BulletBody", parent=body_style, leftIndent=12)

    story = [Paragraph(_clean_html_text(plan_title), title_style), Spacer(1, 12)]
    soup = BeautifulSoup(html, "html.parser")
    root_nodes = soup.body.contents if soup.body else soup.contents

    for node in root_nodes:
        if isinstance(node, NavigableString):
            cleaned = _clean_html_text(str(node))
            if cleaned:
                story.extend([Paragraph(cleaned, body_style), Spacer(1, 4)])
            continue

        if node.name in {"h1", "h2", "h3"}:
            style_name = {"h1": "Heading1", "h2": "Heading2", "h3": "Heading3"}[node.name]
            story.extend([Paragraph(_clean_html_text(node.get_text(" ")), styles[style_name]), Spacer(1, 6)])
            continue

        if node.name in {"ul", "ol"}:
            items = []
            for item in node.find_all("li", recursive=False):
                markup = "".join(_html_to_reportlab_markup(child) for child in item.children)
                items.append(ListItem(Paragraph(markup, bullet_style)))
            story.extend([ListFlowable(items, bulletType="bullet" if node.name == "ul" else "1"), Spacer(1, 8)])
            continue

        markup = "".join(_html_to_reportlab_markup(child) for child in node.children)
        if markup.strip():
            story.extend([Paragraph(markup, body_style), Spacer(1, 4)])

    document.build(story)
    buffer.seek(0)
    return buffer
